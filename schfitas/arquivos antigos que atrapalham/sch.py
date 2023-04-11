import pandas as pd
from pandas import map
import docx 
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL



trashdata = {'Rebarba':0}

pd.set_option("display.max_rows", 250)

tablePD = pd.read_excel("schfitas/400 - 30 dias.xls", usecols=[0, 1, 2] )
tablePd = tablePD.fillna("")
tablePd = tablePD["Unnamed: 1", "Unnamed: 2", "CONTROLE DE ENTREGA DE PEÇAS EM FIBRA DE VIDRO"].str.replace('REBARBA', '', n=1)
tablePD

print(tablePD.iat[2, 2])

doc = docx.Document() 
infoLinha = (tablePD.iat[2, 2])
infoForma = (tablePD.iat[4, 1])
infoCód = (tablePD.iat[4, 0])


# pL = doc.add_paragraph('')
# pL.add_run('Linha:  ').bold = True
# pF = doc.add_paragraph('')
# pF.add_run('Forma:  ').italic = True
# pC = doc.add_paragraph('')
# pC.add_run('Cód:  ').Bold = True


data = (
    ( 'Linha:  ' + infoLinha).bold, 
    ('Forma:  ' + infoForma), 
    ('Cód:  ' + infoCód) 
)
table = doc.add_table(rows=0, cols=1) 
row = table.rows[0].cells 

# row[0].text = 'Linha' + infoLinha

for id in data: 
  
    row = table.add_row().cells 
    
    row[0].text = id
table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
doc.save('gfg.docx')

