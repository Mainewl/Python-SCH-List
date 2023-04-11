from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import pyautogui as pya
import time

pya.PAUSE = 0.15

# lendo o excel archive
wb = load_workbook(r'C:\Users\yury.sell\Desktop\.Net\schfitas\400 - 30 dias.xltx')
sheet = wb.active

# def do word doc
document = Document()

# table
table = document.add_table(rows=3, cols=2)

#img
cell = table.cell(0, 0)
cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
cell.width = Inches(1.25)
paragraph = cell.paragraphs[0]
run = paragraph.add_run()
run.add_picture(r'C:\Users\yury.sell\Desktop\.Net\schfitas\1LOGO - 1 NOVO FORMATO.PNG',width=Inches(1))

#insert de dados
for i in range(1, sheet.max_row + 1):
        #linha nova 
    if i != 1:
        row = table.add_row()
    else:
        row = table.rows[i-1] 

    #texto pra segunda e terceira linha
    cell = row.cells[1]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].add_run(sheet.cell(i, 1).value)

    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].add_run(sheet.cell(i, 2).value)

# salva o doc word
document.save('tabx.docx')

pya.hotkey('win', 'r')
pya.write(r'C:\Users\yury.sell\Desktop\.Net\tabx.docx')
pya.press('enter')
time.sleep(1)

for i in range(4):
    pya.press('alt')
    pya.press('c')
    pya.hotkey('f', 's')
    pya.press('z')
    
    pya.write('semana')
    pya.press('alt')
    pya.hotkey('j', 'l')
    pya.press('x')
    pya.press('l')
while True:
    try:
        location = pya.locateOnScreen(r'C:\Users\yury.sell\Desktop\.Net\schfitas\1close.png', grayscale=True)
        if location == None:
            print("not found yet")
        elif location is not None:
            pya.click(location.left + location.width/2, location.top + location.height/2)
            break
    except Exception as e:
        print(str(e))
    time.sleep(0) 

pya.press('pagedown', presses=50)

for i in range(1, sheet.max_row - 1):
    pya.press('pagueup')


# while True:
#     try:
#         location = pya.locateOnScreen(r'C:\Users\yury.sell\Desktop\.Net\schfitas\SubstituirAll.png', grayscale=True)
#         if location == None:
#             print("not found yet")
#         elif location is not None:
#             pya.click(location.left + location.width/2, location.top + location.height/2)
#             break
#     except Exception as e:
#         print(str(e))
#     time.sleep(0.1)












