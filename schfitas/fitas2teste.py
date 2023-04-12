from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import pyautogui as pya
import time

pya.PAUSE = 0.15
# defs

# defs

# lendo o excel archive
wb = load_workbook(r'C:\Users\yury.sell\Desktop\.Net\schfitas\400 - 30 dias.xltx')
sheet = wb.active
# lendo o excel archive

# def do word doc
document = Document()
# def do word doc

# table
table = document.add_table(rows=3, cols=2)
# table

#img
cell = table.cell(0, 0)
cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
cell.width = Inches(1.25)
paragraph = cell.paragraphs[0]
run = paragraph.add_run()
run.add_picture(r'C:\Users\yury.sell\Desktop\.Net\schfitas\1LOGO - 1 NOVO FORMATO.PNG',width=Inches(1))
#img

#insert de dados
for i in range(1, sheet.max_row + 1):
        #linha nova 
    if i != 1:
        row = table.add_row()
    else:
        row = table.rows[i-1] 

    #texto pra segunda e terceira linha
    cell = row.cells[1]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].add_run(sheet.cell(i, 1).value)

    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].add_run(sheet.cell(i, 2).value)

# salva o doc word
document.save('tab.docx')
# salva o doc word

pya.hotkey('win', 'r')
pya.write(r'C:\Users\yury.sell\Desktop\.Net\tab.docx')
pya.press('enter')
pya.sleep(1)

# tirando as primeiras linhas
pya.press('left')
for i in range(7):
    pya.keyDown('alt')
    pya.hotkey('j', 'l')
    pya.press('x')
    pya.press('l')
    pya.keyUp('alt')
time.sleep(0.1)
# tirando as primeiras linhas

# size
pya.keyDown('alt')
pya.sleep(1)
pya.hotkey('j', 'l')
pya.press('r')
pya.press('s')
pya.keyUp('alt')

pya.keyDown('alt')
pya.sleep(1)
pya.hotkey('c')
pya.hotkey('f', 'y')
pya.keyUp('alt')
pya.write('36')
pya.press('enter')
# size
pya.press('right')
# size
pya.keyDown('alt')
pya.sleep(1)
pya.hotkey('j', 'l')
pya.press('r')
pya.press('s')
pya.keyUp('alt')

pya.keyDown('alt')
pya.sleep(1)
pya.hotkey('c')
pya.hotkey('f', 'y')
pya.keyUp('alt')
pya.write('25')
pya.press('enter')
# size

# retira as "semanas"
for i in range(2):
    pya.keyDown('alt')
    pya.sleep(1)
    pya.press('c')
    pya.hotkey('f', 's')
    pya.press('z')
    pya.keyUp('alt')
    pya.write('semana')

    pya.press('alt')
    pya.sleep(1)
    pya.hotkey('j', 'l')
    pya.press('x')
    pya.press('l')
while True:
    try:
        location = pya.locateOnScreen(r'C:\Users\yury.sell\Desktop\.Net\schfitas\1close.png', grayscale=True)
        if location == None:
            print("not found yet")
        elif location is not None:
            pya.click(location.left + location.width/2, location.top + location.height/2)
            break
    except Exception as e:
        print(str(e))
    time.sleep(0)   
# retira as "semanas"

# select cell y remove laragura
pya.keyDown('alt')
pya.sleep(1)
pya.hotkey('j', 'l')
pya.press('r')
pya.press('b')
pya.keyUp('alt')

pya.keyDown('alt')
pya.sleep(1)
pya.hotkey('j', 'l')
pya.hotkey('i', 'l')
pya.keyUp('alt')
pya.write('0')
pya.press('enter')
# select cell y remove laragura

# mudando a folha pra paisagem
while True:
    try:
        pya.keyDown('alt')
        pya.press('q')
        pya.press('o')
        pya.keyUp('alt')
        location = pya.locateOnScreen(r'C:\Users\yury.sell\Desktop\.Net\schfitas\1paisagem.png', grayscale=True)
        if location == None:
            print("not found yet")
        elif location is not None:
            pya.click(location.left + location.width/2, location.top + location.height/2)
            break
        elif location == location:
            break
    except Exception as e:
        print(str(e))
    time.sleep(0.1)
# mudando a folha pra paisagem
# mudando o formato da folha pra a4 
while True:
    try:
        pya.keyDown('alt')
        pya.press('q')
        pya.hotkey('s', 't')
        pya.keyUp('alt')
        location = pya.locateOnScreen(r'C:\Users\yury.sell\Desktop\.Net\schfitas\1folhaa4.png', grayscale=True)
        if location == None:
            print("not found yet")
        elif location is not None:
            pya.click(location.left + location.width/2, location.top + location.height/2)
            break
        elif location == location:
            break
    except Exception as e:
        print(str(e))
# mudando o formato da folha pra a4 

# select cell y remove spacement
pya.keyDown('alt')
pya.hotkey('j', 'l')
pya.press('r')
pya.press('b')
pya.keyUp('alt')

pya.keyDown('alt')
pya.press('c')
pya.hotkey('p', 'e')
pya.press('n')
pya.keyUp('alt')
# select cell y remove spacement

# selecionando a tabela e botando a borda
pya.keyDown('alt')
pya.sleep(1.5)
pya.hotkey('j', 'l')
pya.press('r')
pya.press('b')
pya.keyUp('alt')

pya.keyDown('alt')
pya.sleep(1.5)
pya.hotkey('j', 't')
pya.hotkey('b', '1')
pya.press('t')
pya.keyUp('alt')
# selecionando a tabela e botando a borda

pya.press('pageup')

# mesclando as celulas
# for i in range(1, sheet.max_row + 1):
for i in range(5):
    pya.keyDown('alt')
    pya.hotkey('j', 'l')
    pya.press('r')
    pya.press('c')
    pya.keyUp('alt')
    pya.keyDown('alt')
    pya.hotkey('j', 'l')
    pya.hotkey('m', 'e')
    pya.keyUp('alt')
    pya.press('down')
# mesclando as celulas
# pya.hotkey('ctrl', 'home')

# save
while True:
    try:
        location = pya.locateOnScreen(r'C:\Users\yury.sell\Desktop\.Net\schfitas\1capsavetela.png', grayscale=True)
        if location == None:
            print("not found yet")
        elif location is not None:
            pya.click(location.left + location.width/2, location.top + location.height/2)
            break
    except Exception as e:
        print(str(e))
    time.sleep(0)

pya.doubleClick()








